from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import RGBColor, Inches
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def test_docx(name):
    # 实例化doc对象
    document = Document()

    # 标题
    document.add_heading('这是一级标题')
    document.add_heading('这是二级标题', level=2)
    # 正文
    document.add_paragraph('  我把话放在这里，这游戏铁定不好玩，都是一波病毒式营销带起来的。中国人根本没能力开发这种游戏。不信你就往我duanceng0417的账户里送一份，你看我玩不玩。')
    # 插入图片
    document.add_picture(r'D:\planself\workspace\language_tools\src\geneaitical\draw_picture\picture\一句话概括人生_0.png', width=Inches(4))

    # 插入内置的项目符号
    list_style = document.styles.add_style('ListStyle', 1)
    list_style.paragraph_format.left_indent = Pt(20)  # 设置缩进
    list_style.paragraph_format.space_after = Pt(6)  # 设置段后间距
    list_style.font.name = 'Arial'  # 设置字体
    list_style.font.size = Pt(12)  # 设置字号
    list_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 设置字体颜色

    document.add_paragraph('❤  说的不错', style='ListStyle')
    document.add_picture(r'D:\planself\workspace\language_tools\src\geneaitical\draw_picture\picture\一句话概括人生_0.png',width=Inches(6))
    document.add_paragraph('❤  嗯，是这样', style='ListStyle')
    document.add_picture(r'D:\planself\workspace\language_tools\src\geneaitical\draw_picture\picture\一句话概括人生_0.png',width=Inches(6))
    document.add_paragraph('❤  对，就是这样', style='ListStyle')
    document.add_picture(r'D:\planself\workspace\language_tools\src\geneaitical\draw_picture\picture\一句话概括人生_0.png',width=Inches(6))

    # 保存
    document.save(name)

def generate_word(keywords, start_word, comments, pic_paths, end_word, cloud_pic_path):
    document = Document()
    # 书写前言
    document.add_paragraph("  " + start_word)
    # 插入图片
    list_style = document.styles.add_style('ListStyle', 1)
    list_style.paragraph_format.left_indent = Pt(20)  # 设置缩进
    list_style.paragraph_format.space_after = Pt(6)  # 设置段后间距
    list_style.font.name = 'Arial'  # 设置字体
    list_style.font.size = Pt(13)  # 设置字号
    list_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 设置字体颜色
    print(str(comments))
    for i in range(len(comments)):
        comment = comments[i]
        pic_path = pic_paths[i]
        document.add_paragraph('❤  ' + comment, style='ListStyle')
        document.add_picture(pic_path, width=Inches(6))
    # 插入词云图
    ciyuStr = '';
    document.add_paragraph("  " + ciyuStr)
    # 画词云图
    document.add_picture(cloud_pic_path, width=Inches(6))
    # 书写结尾
    document.add_paragraph("  " + end_word)
    # 保存
    document.save(keywords + '.docx')

if __name__ == '__main__':
    test_docx("test.docx")